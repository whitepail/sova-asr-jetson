import os
import subprocess
import time
import logging
import uuid
import json
from speech_recognizer import SpeechRecognizer
from number_utils.text2numbers import TextToNumbers
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import numpy as np
from multiprocessing.pool import ThreadPool
from requests import request
from datetime import datetime


speech_recognizer = SpeechRecognizer()
text2numbers = TextToNumbers()


def write_highlighted_text(text, words, document, threshold=70):
    punctuation = ".,?"
    numbers = "0123456789"
    tokens = text.split(" ")
    words_position = 0

    paragraph = document.add_paragraph()

    for token in tokens:
        if len(token) == 0:
            continue

        if token[-1] in punctuation:
            punct = token[-1]
            token = token[:-1]
        else:
            token = token
            punct = ""
        
        if token[0] in numbers:
            paragraph.add_run(token)
        else:
            while words_position < len(words) and words[words_position]["word"] != token.lower():
                words_position += 1

            confidence = words[words_position]["confidence"]

            if confidence < threshold:
                font = paragraph.add_run(token).font
                font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                paragraph.add_run(token)
            
        paragraph.add_run(punct)
        if punct == ".":
            paragraph.add_run("\n")
        else:
            paragraph.add_run(" ")


def punctuator_request(text):
    url = "http://localhost:8890/predict"
    data = {
        "text": text
    }

    response = request("POST", url, json=data)

    result = json.loads(response.text)

    return result


class FileHandler:
    @staticmethod
    def get_recognized_text(blob):
        try:
            # filename = str(uuid.uuid4())
            ts = time.time()
            filename = (os.path.splitext(blob.filename)[0]).replace(" ", "_") + str(datetime.utcfromtimestamp(ts).strftime('_%Y-%m-%d_%H:%M:%S'))
            os.makedirs('./records', exist_ok=True)
            new_record_path = os.path.join('./records', filename + '.webm')
            blob.save(new_record_path)
            audio_file = filename + '.wav'
            converted_record_path = FileHandler.convert_to_wav(new_record_path, audio_file)
            response_models_result = FileHandler.get_models_result(converted_record_path)

            document = Document()
            document.add_heading('Протокол конференции', level=1)
            for result in response_models_result:
                text = result.get('text')
                words = result.get('words')
                write_highlighted_text(text, words, document)

            docx_file = filename + '.docx'
            document.save(f'./records/{docx_file}')

            return 0, audio_file, docx_file, response_models_result

        except Exception as e:
            logging.exception(e)
            return 1,  None, None, str(e)

    @staticmethod
    def convert_to_wav(webm_full_filepath, new_filename):
        converted_record_path = os.path.join('./records', new_filename)
        subprocess.call('ffmpeg -i {0} -ar 16000 -b:a 256k -ac 1 -sample_fmt s16 {1}'.format(
                webm_full_filepath, converted_record_path
            ),
            shell=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        os.remove(webm_full_filepath)
        return converted_record_path

    @staticmethod
    def check_format(files):
        return (files.mimetype.startswith('audio/') or [
            files.filename.endswith(audio_format) for audio_format in [
                'mp3', 'ogg', 'acc', 'flac', 'au', 'm4a', 'mp4', 'mov', 'avi', 'wmv', '3gp', 'flv', 'mkv'
            ]
        ])
        return True

    @staticmethod
    def get_models_result(converted_record_path):
        start = time.time()

        results = []
        decoder_results = speech_recognizer.recognize(converted_record_path)

        score = np.mean([result.get("score") for result in decoder_results])
        words = [w for result in decoder_results for w in result.get("words")]

        texts = [result.get("text") for result in decoder_results]

        texts = [text2numbers.convert(text) for text in [
            " ".join(texts[:len(texts) // 2]),
            " ".join(texts[len(texts) // 2:])
        ]]

        pool = ThreadPool(processes=2)
        texts = pool.map(punctuator_request, texts)
        text = " ".join(texts)

        end = time.time()
        results.append(
            {
                'text': text.strip(),
                'time': round(end - start, 3),
                'confidence': score,
                'words': words
            }
        )
        return results
